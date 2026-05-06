from __future__ import annotations

from pathlib import Path
import json
import re
import shutil
from uuid import uuid4

from fastapi import UploadFile
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever

from .ai_stack import cosine_similarity, dumps_vector, embedding_model, loads_vector
from .config import UPLOAD_DIR
from .db import connect, now_iso, row_to_dict


def chunk_text(text: str, size: int = 900, overlap: int = 120) -> list[str]:
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
        )
        return [chunk.strip() for chunk in splitter.split_text(text) if chunk.strip()]
    except Exception:
        normalized = re.sub(r"\s+", " ", text).strip()
        if not normalized:
            return []
        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            chunks.append(normalized[start : start + size])
            start += max(1, size - overlap)
        return chunks


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
        except Exception:
            return ""
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    return path.read_text(encoding="utf-8", errors="ignore")


def _index_vectors(user_id: int, file_id: int, chunks: list[tuple[int, str]], filename: str) -> None:
    if not chunks:
        return
    vectors = embedding_model().embed_documents([content for _, content in chunks])
    with connect() as db:
        for (chunk_id, content), vector in zip(chunks, vectors):
            metadata = {
                "filename": filename,
                "content_preview": content[:160],
                "vector_store": "sqlite-langchain",
            }
            db.execute(
                """
                INSERT OR REPLACE INTO kb_vectors(chunk_id, file_id, user_id, embedding, metadata, created_at)
                VALUES(?, ?, ?, ?, ?, ?)
                """,
                (chunk_id, file_id, user_id, dumps_vector(vector), json.dumps(metadata, ensure_ascii=False), now_iso()),
            )


def ensure_vector_index() -> int:
    with connect() as db:
        rows = db.execute(
            """
            SELECT c.id AS chunk_id, c.file_id, c.user_id, c.content, f.filename
            FROM kb_chunks c
            JOIN kb_files f ON f.id = c.file_id
            LEFT JOIN kb_vectors v ON v.chunk_id = c.id
            WHERE v.chunk_id IS NULL
            ORDER BY c.id ASC
            """
        ).fetchall()
    grouped: dict[tuple[int, int, str], list[tuple[int, str]]] = {}
    for row in rows:
        key = (row["user_id"], row["file_id"], row["filename"])
        grouped.setdefault(key, []).append((row["chunk_id"], row["content"]))
    for (user_id, file_id, filename), chunks in grouped.items():
        _index_vectors(user_id, file_id, chunks, filename)
    return len(rows)


async def save_upload(user_id: int, upload: UploadFile) -> dict:
    user_dir = UPLOAD_DIR / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{uuid4().hex}_{Path(upload.filename or 'upload.txt').name}"
    stored = user_dir / safe_name
    with stored.open("wb") as out:
        shutil.copyfileobj(upload.file, out)
    text = extract_text(stored)
    chunks = chunk_text(text)
    inserted_chunks: list[tuple[int, str]] = []
    with connect() as db:
        cur = db.execute(
            "INSERT INTO kb_files(user_id, filename, stored_path, created_at) VALUES(?, ?, ?, ?)",
            (user_id, upload.filename or safe_name, str(stored), now_iso()),
        )
        file_id = cur.lastrowid
        for chunk in chunks:
            chunk_cur = db.execute(
                "INSERT INTO kb_chunks(file_id, user_id, content, created_at) VALUES(?, ?, ?, ?)",
                (file_id, user_id, chunk, now_iso()),
            )
            inserted_chunks.append((chunk_cur.lastrowid, chunk))
    _index_vectors(user_id, file_id, inserted_chunks, upload.filename or safe_name)
    return {"id": file_id, "filename": upload.filename, "chunks": len(chunks), "vector_store": "sqlite-langchain"}


def list_files(user_id: int) -> list[dict]:
    with connect() as db:
        rows = db.execute(
            """
            SELECT f.id, f.filename, f.created_at, COUNT(c.id) AS chunks
            FROM kb_files f
            LEFT JOIN kb_chunks c ON c.file_id = f.id
            WHERE f.user_id = ?
            GROUP BY f.id
            ORDER BY f.id DESC
            """,
            (user_id,),
        ).fetchall()
    return [dict(r) for r in rows]


def delete_file(user_id: int, file_id: int) -> None:
    with connect() as db:
        row = row_to_dict(db.execute("SELECT stored_path FROM kb_files WHERE id=? AND user_id=?", (file_id, user_id)).fetchone())
        if row:
            Path(row["stored_path"]).unlink(missing_ok=True)
        db.execute("DELETE FROM kb_files WHERE id=? AND user_id=?", (file_id, user_id))


def _keyword_score(content: str, query: str) -> int:
    terms = [t.lower() for t in re.findall(r"[\w\u4e00-\u9fff]+", query) if len(t) > 1]
    low = content.lower()
    return sum(low.count(term) for term in terms)


def search_documents(user_id: int, query: str, limit: int = 4) -> list[Document]:
    query_vector = embedding_model().embed_query(query or "需求分析")
    with connect() as db:
        rows = db.execute(
            """
            SELECT c.id, c.content, f.filename, v.embedding, v.metadata
            FROM kb_vectors v
            JOIN kb_chunks c ON c.id = v.chunk_id
            JOIN kb_files f ON f.id = c.file_id
            WHERE v.user_id = ?
            ORDER BY c.id DESC
            LIMIT 400
            """,
            (user_id,),
        ).fetchall()
    scored: list[tuple[float, Document]] = []
    for row in rows:
        vector_score = cosine_similarity(query_vector, loads_vector(row["embedding"]))
        lexical_boost = min(_keyword_score(row["content"], query) * 0.05, 0.3)
        metadata = json.loads(row["metadata"] or "{}")
        metadata.update(
            {
                "id": row["id"],
                "filename": row["filename"],
                "score": round(vector_score + lexical_boost, 4),
                "retriever": "langchain-sqlite-vector",
            }
        )
        scored.append(
            (
                vector_score + lexical_boost,
                Document(page_content=row["content"], metadata=metadata),
            )
        )
    return [doc for _, doc in sorted(scored, key=lambda item: item[0], reverse=True)[:limit]]


class KnowledgeVectorRetriever(BaseRetriever):
    user_id: int
    limit: int = 4

    def _get_relevant_documents(self, query: str, *, run_manager) -> list[Document]:
        return search_documents(self.user_id, query, self.limit)


def retriever_for_user(user_id: int, limit: int = 4) -> KnowledgeVectorRetriever:
    return KnowledgeVectorRetriever(user_id=user_id, limit=limit)


def search_knowledge(user_id: int, query: str, limit: int = 4) -> list[dict]:
    documents = retriever_for_user(user_id, limit).invoke(query)
    return [
        {
            "id": int(doc.metadata.get("id", 0)),
            "filename": str(doc.metadata.get("filename", "")),
            "content": doc.page_content,
            "score": float(doc.metadata.get("score", 0.0)),
        }
        for doc in documents
    ]


def format_documents(documents: list[Document]) -> str:
    if not documents:
        return "没有检索到可引用的知识库片段。"
    return "\n\n".join(
        f"[{index}] 来源：{doc.metadata.get('filename', 'unknown')}\n{doc.page_content[:800]}"
        for index, doc in enumerate(documents, start=1)
    )
